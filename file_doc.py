from docx import Document
import aspose.words as aw
from dotenv import load_dotenv
from utils import format_separator
import os

load_dotenv()

path_file_doc = os.getenv("PATH_FILE_DOC", './doc/RETRAITE_PRESTIGE.docx')  # Chemin du fichier Word
pdf_file_path = os.getenv("PATH_FILE_PDF", './pdf/retraite_prestige.pdf')  # Chemin pour sauvegarder le PDF


def convert(duree1, duree2, versement, vers_mens, cotis_total_one,cotis_total_two , cap_acquis_one,cap_acquis_two, plus_value_one, plus_value_two):
    try:
        # Vérifier que le fichier Word existe
        if not os.path.exists(path_file_doc):
            print(f"Erreur : Le fichier Word '{path_file_doc}' est introuvable.")
            return None

        # Charger le document Word
        doc = Document(path_file_doc)

        # Parcourir les paragraphes et remplacer les placeholders par les valeurs fournies
        for paragraph in doc.paragraphs:
            if "{{duree1}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{duree1}}", str(duree1))
            if "{{duree2}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{duree2}}", str(duree2))
            if "{{versement}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{versement}}", format_separator(versement))
            if "{{vers_mens}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{vers_mens}}", format_separator(vers_mens))

            if "{{cotis_total_one}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{cotis_total_one}}", format_separator(cotis_total_one))

            if "{{cotis_total_two}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{cotis_total_two}}", format_separator(cotis_total_two))


            if "{{cap_acquis_one}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{cap_acquis_one}}", format_separator(cap_acquis_one))

            if "{{cap_acquis_two}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{cap_acquis_two}}", format_separator(cap_acquis_two))

            if "{{plus_value_one}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{plus_value_one}}", format_separator(plus_value_one))

            if "{{plus_value_two}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{plus_value_two}}", format_separator(plus_value_two))


        # Sauvegarder le document modifié sous un nouveau nom
        output_word_file = './doc/retraite_prestige_modifie.docx'
        doc.save(output_word_file)

        # Conversion du fichier Word modifié en PDF avec aspose.words
        try:
            doc_aspose = aw.Document(output_word_file)  # Charger le document modifié avec Aspose
            doc_aspose.save(pdf_file_path)  # Sauvegarder en tant que PDF
        except Exception as e:
            print(f"Erreur lors de la conversion en PDF : {e}")
            return None

    except Exception as e:
        print(f"Erreur lors de la lecture du document : {e}")
        return None

    return pdf_file_path  # Retourner le chemin du PDF généré